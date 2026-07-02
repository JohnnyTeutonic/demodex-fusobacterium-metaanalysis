"""
Post-process the pandoc-generated .docx so the reference list renders as a clean,
numbered list. Pandoc flattens a manual LaTeX `thebibliography` into a single
run-on paragraph (and leaks the `{99}` width argument); this script replaces that
paragraph with one numbered paragraph per reference, italicising journal names and
species/genus tokens.

The reference list and order are kept in sync with fusobacterium_demodex_metaanalysis.tex
(ASM citation-sequence order).

Usage:
  python fix_docx_refs.py [path/to/manuscript.docx]
  python fix_docx_refs.py --self-test
"""

from __future__ import annotations

import os
import sys

import docx

HERE = os.path.dirname(__file__)
DEFAULT_DOCX = os.path.join(HERE, os.pardir, "fusobacterium_demodex_metaanalysis.docx")

# (plain text, [substrings to italicise]) -- citation-sequence order.
REFERENCES = [
    ("Chang Y-S, Huang Y-C. Role of Demodex mite infestation in rosacea: a "
     "systematic review and meta-analysis. J Am Acad Dermatol. 2017;77(3):441-447.",
     ["Demodex", "J Am Acad Dermatol"]),
    ("Forton FMN. The pathogenic role of Demodex mites in rosacea: a potential "
     "therapeutic target already in erythematotelangiectatic rosacea? "
     "Dermatol Ther (Heidelb). 2020;10(6):1229-1253.",
     ["Demodex", "Dermatol Ther (Heidelb)"]),
    ("Yamasaki K, Di Nardo A, Bardan A, et al. Increased serine protease activity "
     "and cathelicidin promotes skin inflammation in rosacea. Nat Med. "
     "2007;13(8):975-980.",
     ["Nat Med"]),
    ("Stein Gold L, Kircik L, Fowler J, et al. Efficacy and safety of ivermectin "
     "1% cream in papulopustular rosacea. J Drugs Dermatol. 2014;13(3):316-323.",
     ["J Drugs Dermatol"]),
    ("Taieb A, Ortonne JP, Ruzicka T, et al. Superiority of ivermectin 1% cream "
     "over metronidazole 0.75% cream in treating papulopustular rosacea. "
     "Br J Dermatol. 2015;172(4):1103-1110.",
     ["Br J Dermatol"]),
    ("Nakatsuji T, Cheng JY, Butcher A, et al. Topical ivermectin treatment of "
     "rosacea changes the bacterial microbiome of the skin. J Invest Dermatol. "
     "2025;145(5). doi:10.1016/j.jid.2024.10.592.",
     ["J Invest Dermatol"]),
    ("Olah P, Reuvers N, Radai Z, et al. Microbe-host interaction in rosacea and "
     "its modulation through topical ivermectin. J Invest Dermatol. "
     "2025;145(10):2576-2587.e8.",
     ["J Invest Dermatol"]),
    ("Lacey N, Delaney S, Kavanagh K, Powell FC. Mite-related bacterial antigens "
     "stimulate inflammatory cells in rosacea. Br J Dermatol. 2007;157(3):474-481.",
     ["Br J Dermatol"]),
    ("Murillo N, Aubert J, Raoult D. Microbiota of Demodex mites from rosacea "
     "patients and controls. Microb Pathog. 2014;71-72:37-40.",
     ["Demodex", "Microb Pathog"]),
    ("Rainer BM, Thompson KG, Antonescu C, et al. Characterization and analysis "
     "of the skin microbiota in rosacea: a case-control study. Am J Clin Dermatol. "
     "2020;21(1):139-147.",
     ["Am J Clin Dermatol"]),
    ("Woo YR, Lee SH, Cho SH, Lee JD, Kim HS. Characterization and analysis of the "
     "skin microbiota in rosacea: impact of systemic antibiotics. J Clin Med. "
     "2020;9(1):185.",
     ["J Clin Med"]),
    ("Callahan BJ, McMurdie PJ, Rosen MJ, Han AW, Johnson AJA, Holmes SP. DADA2: "
     "high-resolution sample inference from Illumina amplicon data. Nat Methods. "
     "2016;13(7):581-583.",
     ["Nat Methods"]),
    ("Martin M. Cutadapt removes adapter sequences from high-throughput sequencing "
     "reads. EMBnet J. 2011;17(1):10-12.",
     ["EMBnet J"]),
    ("Quast C, Pruesse E, Yilmaz P, et al. The SILVA ribosomal RNA gene database "
     "project: improved data processing and web-based tools. Nucleic Acids Res. "
     "2013;41(D1):D590-D596.",
     ["Nucleic Acids Res"]),
    ("Callahan BJ, McMurdie PJ, Holmes SP. Exact sequence variants should replace "
     "operational taxonomic units in marker-gene data analysis. ISME J. "
     "2017;11(12):2639-2643.",
     ["ISME J"]),
]


def italic_spans(text: str, tokens):
    """Return list of (substring, is_italic) covering text in order."""
    spans = []
    for tok in tokens:
        i = text.find(tok)
        if i >= 0:
            spans.append((i, i + len(tok)))
    spans.sort()
    # merge/skip overlaps
    merged = []
    for s, e in spans:
        if merged and s < merged[-1][1]:
            continue
        merged.append((s, e))
    out, pos = [], 0
    for s, e in merged:
        if s > pos:
            out.append((text[pos:s], False))
        out.append((text[s:e], True))
        pos = e
    if pos < len(text):
        out.append((text[pos:], False))
    return out


def _find_ref_paragraph(paras):
    for i, p in enumerate(paras):
        t = p.text
        if "Chang Y-S" in t and "mite infestation" in t:
            return i
    return None


def _remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def fix(docx_path: str) -> None:
    d = docx.Document(docx_path)
    idx = _find_ref_paragraph(d.paragraphs)
    if idx is None:
        raise RuntimeError("could not locate the references paragraph")
    # Reuse an existing section heading's style (pandoc docx has no addressable
    # "Heading 1" by name, so copy the style object from another heading).
    heading_style = None
    for p in d.paragraphs:
        if p.text.strip() == "Data availability":
            heading_style = p.style
            break
    # Remove the flattened references paragraph and anything after it.
    for p in d.paragraphs[idx:]:
        _remove_paragraph(p)
    # Rebuild: heading + numbered list.
    head = d.add_paragraph("References")
    if heading_style is not None:
        head.style = heading_style
    for n, (text, toks) in enumerate(REFERENCES, start=1):
        para = d.add_paragraph()
        para.add_run(f"{n}. ")
        for sub, ital in italic_spans(text, toks):
            run = para.add_run(sub)
            run.italic = ital
    d.save(docx_path)
    print(f"Rewrote {len(REFERENCES)} references in {os.path.basename(docx_path)}")


def _self_test() -> None:
    spans = italic_spans("Role of Demodex in J Am Acad Dermatol.",
                         ["Demodex", "J Am Acad Dermatol"])
    joined = "".join(s for s, _ in spans)
    assert joined == "Role of Demodex in J Am Acad Dermatol."
    ital = [s for s, f in spans if f]
    assert "Demodex" in ital and "J Am Acad Dermatol" in ital
    assert len(REFERENCES) == 15
    print("self-test OK: italic_spans + reference count pass.")


def main() -> None:
    if len(sys.argv) > 1 and sys.argv[1] == "--self-test":
        _self_test()
        return
    path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_DOCX
    fix(path)


if __name__ == "__main__":
    main()
