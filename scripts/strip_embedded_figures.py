#!/usr/bin/env python3
"""Produce an mSystems-ready manuscript .docx: remove embedded figure images
but keep their legends (labelled "Figure 1.", "Figure 2." ...), so figures can
be supplied as separate files while the legends stay in the manuscript text.

Usage:
    python3 strip_embedded_figures.py IN.docx OUT.docx
"""
import sys
import docx
from docx.oxml.ns import qn


def has_drawing(p):
    xml = p._p.xml
    return "<w:drawing" in xml or "<pic:pic" in xml


def prepend_bold_label(p, label):
    """Insert a bold run with `label` at the start of paragraph `p`."""
    run = p.add_run(label)
    run.bold = True
    p._p.remove(run._r)
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(run._r)
    else:
        p._p.insert(0, run._r)


def main(in_path, out_path):
    doc = docx.Document(in_path)

    # 1. Remove every paragraph that carries an embedded image.
    removed = 0
    for p in list(doc.paragraphs):
        if has_drawing(p):
            p._p.getparent().remove(p._p)
            removed += 1

    # 2. Drop image relationships so no dangling references remain (otherwise
    #    Word may report the file as corrupt).
    part = doc.part
    dropped = 0
    for rId, rel in list(part.rels.items()):
        if "image" in rel.reltype:
            part.drop_rel(rId)
            dropped += 1

    # 3. Label the surviving caption paragraphs in document order.
    fig_no = 0
    for p in doc.paragraphs:
        if p.style.name == "Image Caption" and not p.text.lstrip().startswith("Figure"):
            fig_no += 1
            prepend_bold_label(p, f"Figure {fig_no}. ")

    doc.save(out_path)
    print(f"removed {removed} embedded image paragraph(s); "
          f"dropped {dropped} image relationship(s); "
          f"labelled {fig_no} figure legend(s); wrote {out_path}")


def _self_test():
    import io
    d = docx.Document()
    d.add_paragraph("body")
    cap = d.add_paragraph("ROC curves ...")
    cap.style = d.styles["Caption"] if "Caption" in [s.name for s in d.styles] else cap.style
    # The self-test only exercises label prepending logic.
    prepend_bold_label(cap, "Figure 1. ")
    assert cap.text.startswith("Figure 1. "), cap.text
    assert cap.runs[0].bold is True
    print("self-test OK")


if __name__ == "__main__":
    if len(sys.argv) == 2 and sys.argv[1] == "--self-test":
        _self_test()
    elif len(sys.argv) == 3:
        main(sys.argv[1], sys.argv[2])
    else:
        sys.exit(__doc__)
